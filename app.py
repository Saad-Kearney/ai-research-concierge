import streamlit as st
import os
import requests
from dotenv import load_dotenv
from docx import Document
from io import BytesIO

# Load environment variables
load_dotenv()
together_api_key = os.getenv("TOGETHER_API_KEY") or st.secrets["TOGETHER_API_KEY"]

API_URL = "https://api.together.xyz/v1/chat/completions"
HEADERS = {
    "Authorization": f"Bearer {together_api_key}",
    "Content-Type": "application/json"
}

st.set_page_config(page_title="ARC – AI Research Concierge", layout="wide")

# Header
col1, col2 = st.columns([5, 1])
with col1:
    st.markdown(
        """
        <h1 style='color:#7823DC;'>Welcome to ARC – AI Research Concierge</h1>
        <h4 style='color:#7823DC; margin-top:-10px;'>by KEARNEY</h4>
        <p style='font-size:16px;'>
            <i>Your personal AI assistant for generating consulting-grade slide structures and Point-of-View (PoV) or Thought Leadership outlines – faster, smarter, and more structured.</i>
        </p>
        """,
        unsafe_allow_html=True
    )
with col2:
    st.image("kearney_logo.png", width=200)

# Goal selection
goal = st.radio(
    "What would you like to generate?",
    ["PoV Structure", "Thought Leadership Structure"],
    index=0
)

# Topic input + model selector
col1, col2 = st.columns([4, 1])
with col1:
    if "topic" not in st.session_state:
        st.session_state.topic = ""

    if st.session_state.topic == "":
        st.markdown("<p style='font-weight:bold; font-size:18px;'>Enter the topic</p>", unsafe_allow_html=True)
        st.session_state.topic = st.text_input("", placeholder="e.g., Future of AI in retail banking")
    else:
        st.text_input("Enter the topic", value=st.session_state.topic, disabled=True)

with col2:
    model_option = st.selectbox("Choose model", [
        "mistralai/Mistral-7B-Instruct-v0.2",
        "mistralai/Mixtral-8x7B-Instruct-v0.1",
        "togethercomputer/llama-2-70b-chat"
    ], index=0)

MODEL = model_option

# Generate
if st.session_state.topic and "summary" not in st.session_state and st.button("Generate Insights"):
    with st.spinner("🧠 Your Kearney AI is thinking..."):

        # Suggested Sources
        prompt_sources = f"List 10 reliable source names or publication types (no links) to study the topic: {st.session_state.topic}"
        sources_body = {
            "model": MODEL,
            "max_tokens": 1000,
            "temperature": 0.3,
            "messages": [{"role": "user", "content": prompt_sources}]
        }
        r_sources = requests.post(API_URL, headers=HEADERS, json=sources_body)
        st.session_state.sources = r_sources.json()['choices'][0]['message']['content'].strip()

        # Summary
        prompt_summary = f"Summarize the key insights about: {st.session_state.topic} for a consulting analyst."
        summary_body = {
            "model": MODEL,
            "max_tokens": 1000,
            "temperature": 0.3,
            "messages": [{"role": "user", "content": prompt_summary}]
        }
        r_summary = requests.post(API_URL, headers=HEADERS, json=summary_body)
        st.session_state.summary = r_summary.json()['choices'][0]['message']['content'].strip()

        # PoV or Thought Leadership
        if goal == "PoV Structure":
            prompt_pov = f"""
You are a strategy consultant at Kearney creating a slide-based Point-of-View on the topic: "{st.session_state.topic}".

Suggest a compelling PoV title first.

Follow this 8-part structure and include slide numbers:
1. Executive Context / Why this matters
2. Key Challenges / Gaps
3. Market Forces / Drivers of Change
4. Strategic Options / Pathways
5. Kearney's Point of View / Recommendation
6. How to Make It Happen / Enablers
7. Next Steps / Roadmap
8. Why Kearney (credentials, experience, assets)

For each slide:
- Start with: Slide [#]: [Title]
- Include 3–5 bullet points
- Suggest sources by name only (e.g., McKinsey, IEA, OECD)

Use concise consulting language. Output in markdown format.
"""
        else:
            prompt_pov = f"""
You are a strategy consultant at Kearney writing a thought leadership article on the topic: "{st.session_state.topic}".

Suggest a compelling article title first.

Follow this 7-part structure:
1. Executive Summary / Why this matters
2. Sector Context / What’s happening now
3. Deeper Diagnostic / What’s broken
4. What Good Looks Like / The Opportunity
5. Strategic Levers / What to Do
6. Implications for Leaders / Who should act
7. Conclusion / Call to Action

Use clear, persuasive business writing. Be factual, bold, and structured.
Output in markdown format.
"""

        pov_body = {
            "model": MODEL,
            "max_tokens": 1200,
            "temperature": 0.3,
            "messages": [{"role": "user", "content": prompt_pov}]
        }
        r_pov = requests.post(API_URL, headers=HEADERS, json=pov_body)
        st.session_state.pov = r_pov.json()['choices'][0]['message']['content'].strip()

# Display
if "summary" in st.session_state and "sources" in st.session_state and "pov" in st.session_state:
    st.markdown("<h2 style='color:#7823DC;'>🔗 Suggested Sources</h2>", unsafe_allow_html=True)
    st.write(st.session_state.sources)

    st.markdown("<h2 style='color:#7823DC;'>📄 Summary</h2>", unsafe_allow_html=True)
    st.write(st.session_state.summary)

    st.markdown(f"<h2 style='color:#7823DC;'>📊 {goal}</h2>", unsafe_allow_html=True)
    st.write(st.session_state.pov)

    # Export to Word
    st.markdown("<h2 style='color:#7823DC;'>⬇️ Export Report</h2>", unsafe_allow_html=True)
    doc = Document()
    doc.add_heading("AI Research Concierge Report", level=1)
    doc.add_paragraph(f"Topic: {st.session_state.topic}")
    doc.add_heading("Suggested Sources", level=2)
    doc.add_paragraph(st.session_state.sources)
    doc.add_heading("Summary", level=2)
    doc.add_paragraph(st.session_state.summary)
    doc.add_heading(goal, level=2)
    doc.add_paragraph(st.session_state.pov)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    st.download_button(
        label="📄 Download as Word (.docx)",
        data=buffer,
        file_name=f"AI_Research_{st.session_state.topic.replace(' ', '_')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    if st.button("🔄 Explore another topic"):
        for key in ["topic", "summary", "sources", "pov"]:
            st.session_state.pop(key, None)
        st.experimental_rerun()
